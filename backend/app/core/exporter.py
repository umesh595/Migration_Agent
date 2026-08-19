"""Renders the 10-deliverable package from the typed MigrationPlan/ArchitectureModel
— never re-generates content as fresh prose (technique #12, DECISIONS.md). Markdown,
Mermaid, and DOCX outputs are all pure functions of the canonical schemas.

Security note (flagged in the PRD review, addressed here): all user/LLM-derived text
that ends up in a Mermaid diagram or a DOCX is sanitized first. Mermaid labels are
attacker-influenced strings (component names, descriptions come from LLM-ingested
free text) rendered client-side by mermaid.js — an unescaped `<script>` or stray
quote/bracket can break the diagram grammar or, worse, get interpreted as HTML by a
loosely-configured renderer. DOCX text runs via python-docx are inserted as literal
text (not interpreted as markup), so injection risk there is inherently low, but
control characters are still stripped for hygiene.
"""

from __future__ import annotations

import io
import re

from docx import Document
from docx.shared import Pt

from app.schemas.architecture import ArchitectureModel
from app.schemas.migration_context import MigrationContext
from app.schemas.migration_plan import MigrationPlan

_CONTROL_CHARS_RE = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f]")


def sanitize_text(value: str) -> str:
    """Strips control characters. Safe baseline for any user/LLM-derived text
    before it's written into any export format."""

    return _CONTROL_CHARS_RE.sub("", value)


def sanitize_mermaid_label(value: str) -> str:
    """Mermaid node labels are wrapped in quotes inside `id["label"]` syntax.
    Strip characters that either break that grammar or could be interpreted as
    markup by a permissive mermaid/HTML renderer downstream."""

    cleaned = sanitize_text(value)
    cleaned = re.sub(r"[<>\"\[\]{}|`]", "", cleaned)
    cleaned = cleaned.replace("\n", " ").strip()
    return cleaned[:120] if len(cleaned) > 120 else cleaned


def sanitize_mermaid_id(component_id: str) -> str:
    """Mermaid node IDs must be alphanumeric/underscore-safe; component ids are
    user-influenced slugs so don't trust them verbatim."""

    return re.sub(r"[^A-Za-z0-9_]", "_", component_id) or "node"


def generate_architecture_mermaid(model: ArchitectureModel) -> str:
    lines = ["graph LR"]
    for component in model.components:
        node_id = sanitize_mermaid_id(component.id)
        label = sanitize_mermaid_label(f"{component.name} ({component.workload_type})")
        lines.append(f'    {node_id}["{label}"]')
    for dep in model.dependencies:
        source = sanitize_mermaid_id(dep.source_id)
        target = sanitize_mermaid_id(dep.target_id)
        label = sanitize_mermaid_label(dep.kind)
        lines.append(f"    {source} -->|{label}| {target}")
    return "\n".join(lines)


def generate_sequence_mermaid(plan: MigrationPlan) -> str:
    lines = ["graph TD"]
    for wave in plan.waves:
        wave_node = f"wave_{wave.index}"
        wave_label = sanitize_mermaid_label(f"Wave {wave.index}")
        lines.append(f'    {wave_node}(["{wave_label}"])')
        for component_id in wave.component_ids:
            node_id = f"{wave_node}_{sanitize_mermaid_id(component_id)}"
            label = sanitize_mermaid_label(component_id)
            lines.append(f'    {node_id}["{label}"]')
            lines.append(f"    {wave_node} --> {node_id}")
    for i in range(len(plan.waves) - 1):
        lines.append(f"    wave_{plan.waves[i].index} --> wave_{plan.waves[i + 1].index}")
    return "\n".join(lines)


def render_markdown(model: ArchitectureModel, plan: MigrationPlan, context: MigrationContext | None) -> str:
    s = sanitize_text
    parts: list[str] = ["# Enterprise Architecture Migration Plan\n"]

    parts.append("## 1. Current Architecture\n")
    parts.append("```mermaid\n" + generate_architecture_mermaid(model) + "\n```\n")
    for c in model.components:
        parts.append(f"- **{s(c.name)}** (`{c.id}`) — {s(c.workload_type)}, {s(c.environment)}"
                      f"{f', {s(c.technology)}' if c.technology else ''}")
    if model.assumptions:
        parts.append("\n**Assumptions:**")
        for a in model.assumptions:
            parts.append(f"- {s(a.text)}")

    parts.append("\n## 2. Target Architecture\n")
    parts.append(s(plan.target_architecture_description))

    parts.append("\n## 3. Component Mapping\n")
    parts.append("| Component | Disposition | Target |\n|---|---|---|")
    for m in plan.component_mappings:
        parts.append(f"| {s(m.component_id)} | {m.disposition} | {s(m.target_description)} |")

    parts.append("\n## 4. Component Migration Approach\n")
    for p in plan.component_plans:
        parts.append(f"### {s(p.component_id)} (wave {p.wave_index}, {p.disposition})")
        for step in p.steps:
            parts.append(f"1. {s(step)}")
        if p.estimated_effort:
            parts.append(f"- Estimated effort: {s(p.estimated_effort)}")

    parts.append("\n## 5. Migration Sequence\n")
    parts.append("```mermaid\n" + generate_sequence_mermaid(plan) + "\n```\n")
    for w in plan.waves:
        parts.append(f"- **Wave {w.index}**: {', '.join(s(c) for c in w.component_ids)} — {s(w.rationale)}")
        for g in w.coexistence_groups:
            parts.append(f"  - _Coexistence ({', '.join(s(c) for c in g.component_ids)})_: {s(g.coexistence_strategy)}")

    parts.append("\n## 6. Risks & Assumptions\n")
    for r in plan.risks:
        parts.append(f"- **[{r.severity}]** {s(r.description)} — _mitigation:_ {s(r.mitigation)}")

    parts.append("\n## 7. Validation Approach\n")
    if plan.validation_summary:
        parts.append(s(plan.validation_summary.overall_strategy))
        for check in plan.validation_summary.cross_component_checks:
            parts.append(f"- {s(check.check_type)}: {s(check.description)}")

    parts.append("\n## 8. Cutover Strategy\n")
    if plan.cutover_strategy:
        parts.append(s(plan.cutover_strategy.approach))
        for step in plan.cutover_strategy.steps:
            parts.append(f"1. {s(step)}")
        parts.append("\n**Go/No-Go criteria:**")
        for c in plan.cutover_strategy.go_no_go_criteria:
            parts.append(f"- {s(c)}")

    parts.append("\n## 9. Rollback Strategy\n")
    if plan.rollback_strategy:
        parts.append(s(plan.rollback_strategy.approach))
        for step in plan.rollback_strategy.steps:
            parts.append(f"1. {s(step)}")

    parts.append("\n## 10. Migration Roadmap\n")
    parts.append(
        "| Wave | Component | Disposition | Summary | Owner | Effort | Depends on waves |\n"
        "|---|---|---|---|---|---|---|"
    )
    for item in plan.roadmap_items:
        parts.append(
            f"| {item.wave_index} | {s(item.component_id)} | {item.disposition} | {s(item.summary)} | "
            f"{s(item.owner_placeholder)} | {s(item.estimated_effort or '-')} | "
            f"{', '.join(str(w) for w in item.depends_on_waves) or '-'} |"
        )

    if context:
        parts.append("\n## Migration Context\n")
        parts.append(f"- Source: {context.source_environment} → Target: {context.target_environment}")
        parts.append(f"- Downtime tolerance: {context.downtime_tolerance}")
        for constraint in context.constraints:
            parts.append(f"- Constraint: {s(constraint)}")

    return "\n".join(parts)


def render_docx(model: ArchitectureModel, plan: MigrationPlan, context: MigrationContext | None) -> bytes:
    s = sanitize_text
    doc = Document()
    doc.styles["Normal"].font.size = Pt(10.5)

    doc.add_heading("Enterprise Architecture Migration Plan", level=0)

    doc.add_heading("1. Current Architecture", level=1)
    for c in model.components:
        doc.add_paragraph(
            f"{s(c.name)} ({c.id}) — {s(c.workload_type)}, {s(c.environment)}"
            f"{f', {s(c.technology)}' if c.technology else ''}",
            style="List Bullet",
        )

    doc.add_heading("2. Target Architecture", level=1)
    doc.add_paragraph(s(plan.target_architecture_description))

    doc.add_heading("3. Component Mapping", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = "Component", "Disposition", "Target"
    for m in plan.component_mappings:
        row = table.add_row().cells
        row[0].text, row[1].text, row[2].text = s(m.component_id), str(m.disposition), s(m.target_description)

    doc.add_heading("4. Component Migration Approach", level=1)
    for p in plan.component_plans:
        doc.add_heading(f"{s(p.component_id)} (wave {p.wave_index}, {p.disposition})", level=2)
        for step in p.steps:
            doc.add_paragraph(s(step), style="List Number")

    doc.add_heading("5. Migration Sequence", level=1)
    for w in plan.waves:
        doc.add_paragraph(f"Wave {w.index}: {', '.join(s(c) for c in w.component_ids)} — {s(w.rationale)}",
                           style="List Bullet")
        for g in w.coexistence_groups:
            doc.add_paragraph(
                f"Coexistence ({', '.join(s(c) for c in g.component_ids)}): {s(g.coexistence_strategy)}",
                style="List Bullet 2",
            )

    doc.add_heading("6. Risks & Assumptions", level=1)
    for r in plan.risks:
        doc.add_paragraph(f"[{r.severity}] {s(r.description)} — mitigation: {s(r.mitigation)}", style="List Bullet")

    doc.add_heading("7. Validation Approach", level=1)
    if plan.validation_summary:
        doc.add_paragraph(s(plan.validation_summary.overall_strategy))
        for check in plan.validation_summary.cross_component_checks:
            doc.add_paragraph(f"{s(check.check_type)}: {s(check.description)}", style="List Bullet")

    doc.add_heading("8. Cutover Strategy", level=1)
    if plan.cutover_strategy:
        doc.add_paragraph(s(plan.cutover_strategy.approach))
        for step in plan.cutover_strategy.steps:
            doc.add_paragraph(s(step), style="List Number")

    doc.add_heading("9. Rollback Strategy", level=1)
    if plan.rollback_strategy:
        doc.add_paragraph(s(plan.rollback_strategy.approach))
        for step in plan.rollback_strategy.steps:
            doc.add_paragraph(s(step), style="List Number")

    doc.add_heading("10. Migration Roadmap", level=1)
    table = doc.add_table(rows=1, cols=6)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, title in enumerate(["Wave", "Component", "Disposition", "Summary", "Owner", "Effort"]):
        hdr[i].text = title
    for item in plan.roadmap_items:
        row = table.add_row().cells
        row[0].text = str(item.wave_index)
        row[1].text = s(item.component_id)
        row[2].text = str(item.disposition)
        row[3].text = s(item.summary)
        row[4].text = s(item.owner_placeholder)
        row[5].text = s(item.estimated_effort or "-")

    if context:
        doc.add_heading("Migration Context", level=1)
        doc.add_paragraph(f"Source: {context.source_environment} -> Target: {context.target_environment}")
        doc.add_paragraph(f"Downtime tolerance: {context.downtime_tolerance}")

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
